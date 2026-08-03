import os
import json
import logging
import warnings
import re
import asyncio  
import html  # Añadido para la Modificación 1 y 3: Escapar caracteres y prevenir fallos de parseo
from datetime import datetime, timedelta  # Añadido timedelta para la Modificación 2: Control de Mute 24h
from dotenv import load_dotenv # type: ignore
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ChatMember, ChatPermissions # type: ignore
from telegram.ext import ( # type: ignore
    Application, 
    CommandHandler, 
    MessageHandler, 
    CallbackQueryHandler, 
    ChatMemberHandler, 
    filters, 
    ContextTypes
)
from google import genai
from google.genai import types # type: ignore

# Librerías añadidas para la Opción 3 (Generación de Informes Word Estrictos)
import docx # type: ignore
from docx.shared import Pt, Cm # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH # type: ignore

# Silenciar advertencias estéticas innecesarias en la consola
warnings.filterwarnings("ignore", category=UserWarning)

# CONFIGURACIÓN DE HISTORIAL DE ACCIONES (LOGS)
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', 
    level=logging.INFO
)

# Cargar variables de entorno de forma segura desde el archivo .env
ruta_env = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
load_dotenv(dotenv_path=ruta_env)

TOKEN = os.getenv("TELEGRAM_TOKEN")
GEMINI_KEY = os.getenv("AI_API_KEY")

if not TOKEN or not GEMINI_KEY:
    raise ValueError("ERROR CRÍTICO: Configura correctamente TELEGRAM_TOKEN y AI_API_KEY en tu archivo .env")

# Inicialización oficial del cliente de Gemini (Nueva SDK)
ai_client = genai.Client(api_key=GEMINI_KEY)

# Archivos locales de base de datos JSON
PROF_FILE = "config_profesor.json"
GROUPS_FILE = "grupos.json"
PAGOS_FILE = "pagos.json"
EVENTOS_FILE = "eventos.json" 
CONTRASENA_CORRECTA = "profesor.udomonagas"

# --- NUEVOS ARCHIVOS DE PERSISTENCIA (MODIFICACIONES 2 Y 3) ---
REGLAS_FILE = "reglas.json"
COMPORTAMIENTO_FILE = "comportamiento.json"
CONSULTAS_FILE = "consultas_diarias.json"


# =========================================================================
# GESTIÓN DE PERSISTENCIA (LECTURA Y ESCRITURA EN DISCO)
# =========================================================================

# --- GESTIÓN MULTI-PROFESOR ---
def cargar_profesores():
    if os.path.exists(PROF_FILE):
        try:
            with open(PROF_FILE, "r") as f:
                data = json.load(f)
                return data.get("profesores", [])
        except Exception:
            return []
    return []

def guardar_profesor(user_id):
    profesores = cargar_profesores()
    if user_id not in profesores:
        profesores.append(user_id)
        with open(PROF_FILE, "w") as f:
            json.dump({"profesores": profesores}, f)

def es_profesor_autorizado(user_id):
    return user_id in cargar_profesores()


# --- CONTROL DE GRUPOS VINCULADOS ---
def cargar_grupos():
    if os.path.exists(GROUPS_FILE):
        try:
            with open(GROUPS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def guardar_grupo(chat_id, titulo, id_profesor, total_estudiantes=34):
    grupos = cargar_grupos()
    grupos[str(chat_id)] = {
        "titulo": titulo,
        "id_profesor": id_profesor,
        "total_estudiantes": total_estudiantes 
    }
    with open(GROUPS_FILE, "w", encoding="utf-8") as f:
        json.dump(grupos, f, ensure_ascii=False, indent=4)

def eliminar_grupo(chat_id):
    grupos = cargar_grupos()
    str_chat_id = str(chat_id)
    if str_chat_id in grupos:
        del grupos[str_chat_id]
        with open(GROUPS_FILE, "w", encoding="utf-8") as f:
            json.dump(grupos, f, ensure_ascii=False, indent=4)


# --- CONTROL DE EVENTOS DE RECAUDACIÓN (OPCIÓN 1) ---
def cargar_eventos():
    if os.path.exists(EVENTOS_FILE):
        try:
            with open(EVENTOS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def guardar_evento(chat_id, evento_activo, fecha_limite, estudiantes_objetivo):
    eventos = cargar_eventos()
    eventos[str(chat_id)] = {
        "evento_activo": evento_activo,
        "fecha_limite": fecha_limite,  
        "estudiantes_objetivo": estudiantes_objetivo,
        "estado": "abierto"
    }
    with open(EVENTOS_FILE, "w", encoding="utf-8") as f:
        json.dump(eventos, f, ensure_ascii=False, indent=4)

def obtener_evento_activo(chat_id):
    eventos = cargar_eventos()
    return eventos.get(str(chat_id))

def cerrar_evento_en_json(chat_id):
    eventos = cargar_eventos()
    str_cid = str(chat_id)
    if str_cid in eventos:
        eventos[str_cid]["estado"] = "cerrado"
        with open(EVENTOS_FILE, "w", encoding="utf-8") as f:
            json.dump(eventos, f, ensure_ascii=False, indent=4)


# --- BASE DE DATOS E INGESTA DE PAGOS POR EVENTO ---
def cargar_pagos():
    if os.path.exists(PAGOS_FILE):
        try:
            with open(PAGOS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def registrar_pago_valido(chat_id, estudiante, referencia, evento_nombre):
    pagos = cargar_pagos()
    str_chat_id = str(chat_id)
    
    if str_chat_id not in pagos:
        pagos[str_chat_id] = {}
    if evento_nombre not in pagos[str_chat_id]:
        pagos[str_chat_id][evento_nombre] = []
        
    for pago in pagos[str_chat_id][evento_nombre]:
        if pago["referencia"] == referencia:
            return "duplicado_referencia"
        if p_estudiante := pago.get("estudiante"):
            if p_estudiante.lower() == estudiante.lower():
                return "duplicado_estudiante"
            
    pagos[str_chat_id][evento_nombre].append({
        "estudiante": estudiante,
        "referencia": referencia,
        "fecha_registro": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    })
    
    with open(PAGOS_FILE, "w", encoding="utf-8") as f:
        json.dump(pagos, f, ensure_ascii=False, indent=4)
    return "exito"


# --- PERSISTENCIA EXCLUSIVA PARA LAS MODIFICACIONES 2 Y 3 ---
def cargar_reglas():
    if os.path.exists(REGLAS_FILE):
        try:
            with open(REGLAS_FILE, "r", encoding="utf-8") as f: return json.load(f)
        except Exception: return {}
    return {}

def guardar_reglas(chat_id, texto_reglas):
    reglas = cargar_reglas()
    reglas[str(chat_id)] = texto_reglas
    with open(REGLAS_FILE, "w", encoding="utf-8") as f:
        json.dump(reglas, f, ensure_ascii=False, indent=4)

def cargar_comportamiento():
    if os.path.exists(COMPORTAMIENTO_FILE):
        try:
            with open(COMPORTAMIENTO_FILE, "r", encoding="utf-8") as f: return json.load(f)
        except Exception: return {}
    return {}

def guardar_comportamiento(data):
    with open(COMPORTAMIENTO_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def cargar_consultas():
    if os.path.exists(CONSULTAS_FILE):
        try:
            with open(CONSULTAS_FILE, "r", encoding="utf-8") as f: return json.load(f)
        except Exception: return {}
    return {}

def guardar_consultas(data):
    with open(CONSULTAS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


# =========================================================================
# FUNCIONES AUXILIARES DE CONTEO Y MAQUETADO ACADÉMICO
# =========================================================================

# MODIFICACIÓN 1: FILTRO DE SANITIZACIÓN DE ENTIDADES HTML PARA TELEGRAM
def sanitizar_html_telegram(texto: str) -> str:
    if not texto:
        return ""
    # Traducir variaciones de saltos de línea web a nativos (\n)
    texto = re.sub(r'<br\s*/?>', '\n', texto, flags=re.IGNORECASE)
    # Sanitizar párrafos
    texto = texto.replace('<p>', '').replace('<P>', '')
    texto = re.sub(r'</p\s*>', '\n', texto, flags=re.IGNORECASE)
    # Limpiar contenedores de listas desordenadas web
    texto = texto.replace('<ul>', '').replace('</ul>', '').replace('<UL>', '').replace('</UL>', '')
    texto = re.sub(r'<li\s*>', '• ', texto, flags=re.IGNORECASE)
    texto = re.sub(r'</li\s*>', '\n', texto, flags=re.IGNORECASE)
    return texto.strip()

async def obtener_cantidad_estudiantes(bot, chat_id_grupo: int) -> int:
    try:
        total_miembros = await bot.get_chat_member_count(chat_id_grupo)
        return max(0, total_miembros - 1)
    except Exception as e:
        logging.error(f"Error al contar miembros nativos: {e}")
        return 34

def crear_documento_academico(texto_estructurado_ia, nombre_archivo="Material_de_Estudio.docx"):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)

    lineas = texto_estructurado_ia.split('\n')
    for linea in lineas:
        linea = linea.strip()
        if not linea:
            continue
        
        if linea.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(linea.replace('# ', ''))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
             
        elif linea.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(linea.replace('## ', ''))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(linea)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    ruta_guardado = os.path.join(os.path.dirname(os.path.abspath(__file__)), nombre_archivo)
    doc.save(ruta_guardado)
    return ruta_guardado


# =========================================================================
# HANDLERS Y FLUJO DE COMANDOS / INTERFACES
# =========================================================================

async def start_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_chat.type != "private":
        return

    id_actual = update.effective_user.id
    if es_profesor_autorizado(id_actual):
        await update.message.reply_text("✨ Acceso confirmado. Ya estás autenticado dentro de tu espacio privado del Servidor Centralizado.\nPuedes enviarme un audio, documento o enlace cuando gustes.")
        return

    botones = [[
        InlineKeyboardButton("Sí", callback_data="auth_si"), 
        InlineKeyboardButton("No", callback_data="auth_no")
    ]]
    teclado = InlineKeyboardMarkup(botones)
    await update.message.reply_text(
        "🔒 Este bot es para el uso exclusivo de profesores (Soporte Multi-Profesor Activo).\n\n¿Eres un profesor?",
        reply_markup=teclado
    )

async def iniciar_recaudacion_comando(update: Update, context: ContextTypes.DEFAULT_TYPE):
    id_profesor = update.effective_user.id
    if update.effective_chat.type != "private" or not es_profesor_autorizado(id_profesor):
        return
        
    grupos = cargar_grupos()
    grupos_propios = {cid: info for cid, info in grupos.items() if info.get("id_profesor") == id_profesor}
    
    if not grupos_propios:
        await update.message.reply_text("⚠️ No posees secciones vinculadas para coordinar una recaudación de fondos.")
        return
        
    botones = []
    for cid, info in grupos_propios.items():
        botones.append([InlineKeyboardButton(f"💰 {info['titulo']}", callback_data=f"recaudar_g:{cid}")])
    teclado = InlineKeyboardMarkup(botones)
    
    await update.message.reply_text("🎯 Selecciona el grupo estudiantil en el cual deseas aperturar la recaudación de fondos:", reply_markup=teclado)

# MODIFICACIÓN 2: COMANDO PARA DICTAR EL REGLAMENTO INTERNO
async def establecer_reglas_comando(update: Update, context: ContextTypes.DEFAULT_TYPE):
    id_profesor = update.effective_user.id
    if update.effective_chat.type != "private" or not es_profesor_autorizado(id_profesor):
        return
        
    grupos = cargar_grupos()
    grupos_propios = {cid: info for cid, info in grupos.items() if info.get("id_profesor") == id_profesor}
    
    if not grupos_propios:
        await update.message.reply_text("⚠️ No posees secciones vinculadas para establecer un reglamento.")
        return
        
    botones = []
    for cid, info in grupos_propios.items():
        botones.append([InlineKeyboardButton(f"📜 Reglas: {info['titulo']}", callback_data=f"config_reglas_g:{cid}")])
    teclado = InlineKeyboardMarkup(botones)
    
    await update.message.reply_text("🎯 Selecciona el grupo estudiantil para el cual vas a dictar las reglas de convivencia:", reply_markup=teclado)

async def recibir_documento_profesor(update: Update, context: ContextTypes.DEFAULT_TYPE):
    id_profesor = update.effective_user.id
    if update.effective_chat.type != "private" or not es_profesor_autorizado(id_profesor):
        return
        
    documento = update.message.document
    context.user_data['material_pendiente'] = {
        "tipo": "documento",
        "file_id": documento.file_id,
        "nombre": documento.file_name,
        "caption": update.message.caption or ""
    }
    await mostrar_lista_grupos_material(update, context)

async def mostrar_lista_grupos_material(update: Update, context: ContextTypes.DEFAULT_TYPE):
    id_profesor = update.effective_user.id
    grupos = cargar_grupos()
    grupos_propios = {cid: info for cid, info in grupos.items() if info.get("id_profesor") == id_profesor}
    
    if not grupos_propios:
        await update.message.reply_text("⚠️ No tienes grupos académicos indexados bajo tu control para distribuir este recurso.")
        return
        
    botones = []
    for chat_id, info in grupos_propios.items():
        botones.append([InlineKeyboardButton(f"📁 {info['titulo']}", callback_data=f"env_mat:{chat_id}")])
    botones.append([InlineKeyboardButton("❌ Cancelar Envío", callback_data="cancelar_envio_mat")])
    teclado = InlineKeyboardMarkup(botones)
    
    material = context.user_data.get('material_pendiente')
    nombre_mat = material.get('nombre', 'Enlace / Bibliografía') if material else "Recurso"
    tipo_mat = material.get('tipo', 'UNKN').upper() if material else ""
    
    await update.message.reply_text(
        f"📦 <b>Material de apoyo detectado:</b>\n• <i>{nombre_mat}</i> ({tipo_mat})\n\n"
        f"¿A qué comunidad de alumnos deseas despachar este recurso?",
        reply_markup=teclado,
        parse_mode="HTML"
    )


# =========================================================================
# ORQUESTADOR CENTRAL DE MENSAJES DE TEXTO PRIVADOS
# =========================================================================
async def manejar_mensajes_texto_privado(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_chat.type != "private":
        return

    id_usuario = update.effective_user.id
    texto_ingresado = update.message.text.strip()

    if context.user_data.get('esperando_clave'):
        context.user_data['esperando_clave'] = False
        if texto_ingresado == CONTRASENA_CORRECTA:
            guardar_profesor(id_usuario)
            await update.message.reply_text("🔑 Contraseña correcta. Acceso concedido.\nHas sido añadido exitosamente al entorno distribuido multi-profesor.")
        else:
            await update.message.reply_text("❌ Contraseña incorrecta. Acceso denegado.")
        return

    if not es_profesor_autorizado(id_usuario):
        return

    if context.user_data.get('esperando_recaudacion_datos'):
        chat_id_grupo = context.user_data['esperando_recaudacion_datos']
        if "|" not in texto_ingresado:
            await update.message.reply_text("❌ Formato incorrecto. Debes separar el nombre y la fecha usando el carácter '|'. Introduce los datos nuevamente.")
            return
            
        context.user_data['esperando_recaudacion_datos'] = False
        parts = texto_ingresado.split("|")
        evento_nombre = parts[0].strip()
        fecha_str = parts[1].strip()
        
        try:
            fecha_dt = datetime.strptime(fecha_str, "%Y-%m-%d %H:%M")
            fecha_limite_completa = fecha_dt.strftime("%Y-%m-%d %H:%M:%S")
            
            if datetime.now() > fecha_dt:
                await update.message.reply_text("❌ La fecha límite establecida ya ha expirado en el tiempo real. Ejecuta nuevamente el comando /recaudar.")
                return

            total_alumnos = await obtener_cantidad_estudiantes(context.bot, chat_id_grupo)
            guardar_evento(chat_id_grupo, evento_nombre, fecha_limite_completa, total_alumnos)
            
            context.application.job_queue.run_once(
                ejecutar_cierre_automatico,
                when=fecha_dt,
                data={
                    "chat_id": chat_id_grupo,
                    "evento_nombre": evento_nombre,
                    "id_profesor": id_usuario
                }
            )
            
            grupos = cargar_grupos()
            nombre_g = grupos.get(str(chat_id_grupo), {}).get("titulo", "Grupo")
             
            await update.message.reply_text(
                f"✅ <b>Recaudación Abierta Exitosamente</b>\n\n"
                f"• <b>Sección:</b> {nombre_g}\n"
                f"• <b>Motivo:</b> {evento_nombre}\n"
                f"• <b>Cierre Plazo:</b> {fecha_limite_completa}\n"
                f"• <b>Meta Estudiantes:</b> {total_alumnos} alumnos.\n\n"
                f"El bot ha tomado control del canal y auditará de forma multimodal los captures de pantalla.",
                parse_mode="HTML"
            )
            
            await context.bot.send_message(
                chat_id=chat_id_grupo,
                text=f"📢 <b>[SISTEMA ACADÉMICO] Control de Recaudación Abierto</b>\n\n"
                     f"El profesor ha abierto el proceso de recaudación para: <b>{evento_nombre}</b>.\n"
                     f"⏰ <b>Fecha y Hora Límite de Pago:</b> <code>{fecha_limite_completa}</code>.\n\n"
                     f"<i>Por favor envíen por esta vía el capture de pantalla nítido de su transferencia bancaria. El Delegado Virtual validará su identidad y número de referencia.</i>",
                parse_mode="HTML"
            )
        except ValueError:
            await update.message.reply_text("❌ Estructura cronológica inválida. Respeta rigurosamente el formato: AAAA-MM-DD HH:MM.")
        return

    patron_url = re.search(r'(https?://[^\s]+)', texto_ingresado)
    if patron_url:
        context.user_data['material_pendiente'] = {
            "tipo": "enlace",
            "contenido": texto_ingresado
        }
        await mostrar_lista_grupos_material(update, context)
        return
        
    await update.message.reply_text("✨ Servidor Central de Profesores Activo. Puedes enviar un archivo, un enlace web, grabar un audio, configurar normas con /establecer_reglas o iniciar una recaudación con /recaudar.")


# =========================================================================
# DETECCIÓN Y ASIGNACIÓN AUTOMÁTICA DE COMUNIDADES
# =========================================================================
async def detectar_nuevo_grupo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    cambio_miembro = update.my_chat_member
    if not cambio_miembro:
        return

    estado_anterior = cambio_miembro.old_chat_member.status
    estado_nuevo = cambio_miembro.new_chat_member.status
    chat = cambio_miembro.chat

    if chat.type in ["group", "supergroup"]:
        if estado_nuevo in [ChatMember.MEMBER, ChatMember.ADMINISTRATOR] and estado_anterior not in [ChatMember.MEMBER, ChatMember.ADMINISTRATOR]:
            invitador = cambio_miembro.from_user
            
            if not invitador or not es_profesor_autorizado(invitador.id):
                logging.warning(f"Bot agregado al grupo '{chat.title}' por usuario no verificado.")
                return

            id_profesor = invitador.id
            context.bot_data[f"titulo_pendiente_{chat.id}"] = chat.title
            
            botones = [[
                InlineKeyboardButton("Aceptar y vincular a mi espacio", callback_data=f"g_aceptar:{chat.id}"),
                InlineKeyboardButton("Rechazar y salir", callback_data=f"g_rechazar:{chat.id}")
            ]]
            teclado = InlineKeyboardMarkup(botones)
            
            await context.bot.send_message(
                chat_id=id_profesor,
                text=f"🚨 <b>Alerta de nuevo grupo detectado</b>\n\nHas añadido al bot al entorno estudiantil: <b>{chat.title}</b>.\n¿Deseas autorizar la indexación en tu cuenta?",
                reply_markup=teclado,
                parse_mode="HTML"
            )


# =========================================================================
# MOTOR DE AUDIO MULTIMODAL CON MATRIZ DE DURACIÓN
# =========================================================================
async def procesar_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_chat.type != "private":
        return

    id_actual = update.effective_user.id
    if not es_profesor_autorizado(id_actual):
        await update.message.reply_text("❌ Acceso denegado. No posees permisos administrativos.")
        return

    audio_obj = update.message.voice or update.message.audio
    duracion_segundos = audio_obj.duration

    # --- MODIFICACIÓN 2: INTERCEPTACIÓN SI ESTÁ CONFIGURANDO REGLAMENTO ---
    chat_id_reglas = context.user_data.get('esperando_audio_reglas')
    if chat_id_reglas:
        await update.message.reply_text("🎙️ Audio de reglamento recibido. Procesando y memorizando normas con Inteligencia Artificial...")
        archivo_voice = await audio_obj.get_file()
        datos_audio_bytes = await archivo_voice.download_as_bytearray()
        context.user_data['esperando_audio_reglas'] = None
        await ejecutar_flujo_configurar_reglas(update, context, datos_audio_bytes, chat_id_reglas)
        return
    # ----------------------------------------------------------------------

    grupos_totales = cargar_grupos()
    grupos_activos = {cid: info["titulo"] for cid, info in grupos_totales.items() if info.get("id_profesor") == id_actual}

    if not grupos_activos:
        await update.message.reply_text("⚠️ No posees grupos académicos validados en tu registro privado. Añade el bot a una sección primero.")
        return

    await update.message.reply_text("🎙️ Audio recibido de forma segura. Descargando búfer de voz...")
    
    archivo_voice = await audio_obj.get_file()
    datos_audio_bytes = await archivo_voice.download_as_bytearray()

    context.user_data['audio_bytes_pendientes'] = list(datos_audio_bytes)

    if duracion_segundos >= 120:
        botones = [
            [InlineKeyboardButton("💬 Opción 1: Transmitir Informativo", callback_data="audio_opt:mensaje")],
            [InlineKeyboardButton("📄 Opción 2: Generar Trabajo Escrito (Word)", callback_data="audio_opt:informe")],
            [InlineKeyboardButton("❌ Cancelar Operación", callback_data="audio_opt:cancelar")]
        ]
        teclado = InlineKeyboardMarkup(botones)
        await update.message.reply_text(
            f"⏳ <b>Nota de voz de larga duración detectada ({duracion_segundos // 60} min {duracion_segundos % 60} seg).</b>\n"
            f"¿Qué acción estructural deseas que ejecute el Delegado Virtual con esta clase magistral?",
            reply_markup=teclado,
            parse_mode="HTML"
        )
    else:
        await ejecutar_flujo_audio_corto(update, context, datos_audio_bytes, grupos_activos)

async def ejecutar_flujo_audio_corto(update_or_query, context, datos_audio_bytes, grupos_activos):
    is_query = hasattr(update_or_query, "edit_message_text")
    
    prompt_instrucciones = (
        "Actúas como un Operador de Minutas Académicas profesional.\n"
        "Analiza detalladamente la nota de voz provista. Tu tarea es extraer de forma estructurada los puntos clave, "
        "tareas asignadas, fechas de entrega u observaciones especiales descritas por el profesor.\n"
        "Reglas críticas:\n"
        "1. Limpia cualquier chiste, muletilla, ruido o comentario informal ajeno a la materia.\n"
        "2. Usa FORMATO HTML EXCLUSIVO admitido por Telegram: <b>texto para negritas</b>, <i>texto para cursivas</i> y viñetas estándar (•).\n"
        "3. Redacta de forma clara, directa y formal."
    )

    max_intentos = 3
    respuesta_ia = None
    
    for intento in range(max_intentos):
        try:
            respuesta_ia = ai_client.models.generate_content(
                model='gemini-2.5-flash', 
                contents=[
                    types.Part.from_bytes(data=bytes(datos_audio_bytes), mime_type="audio/ogg"),
                    prompt_instrucciones
                ]
            )
            break
        except Exception as e:
            if ("503" in str(e) or "UNAVAILABLE" in str(e)) and intento < max_intentos - 1:
                await asyncio.sleep(4)
            else:
                msg = f"❌ Error interno al procesar con Gemini tras reintentos:\n<code>{str(e)}</code>"
                if is_query: await update_or_query.edit_message_text(msg, parse_mode="HTML")
                else: await update_or_query.message.reply_text(msg, parse_mode="HTML")
                return

    if respuesta_ia and respuesta_ia.text:
        # --- MODIFICACIÓN 1: BLINDAJE Y SANITIZACIÓN HTML CONTRA FALLOS DE PARSEO ---
        context.user_data['minuta_pendiente'] = sanitizar_html_telegram(respuesta_ia.text)
        # ---------------------------------------------------------------------------
        
        botones = []
        grupos_sin_permisos = []
        
        for chat_id, nombre_grupo in grupos_activos.items():
            es_admin = True
            try:
                miembro_bot = await context.bot.get_chat_member(chat_id=chat_id, user_id=context.bot.id)
                if miembro_bot.status not in [ChatMember.ADMINISTRATOR, ChatMember.OWNER]:
                    es_admin = False
            except Exception:
                es_admin = False
        
            if not es_admin:
                grupos_sin_permisos.append(nombre_grupo)
                botones.append([InlineKeyboardButton(text=f"⚠️ {nombre_grupo} (Sin Admin)", callback_data=f"send:{chat_id}")])
            else:
                botones.append([InlineKeyboardButton(text=nombre_grupo, callback_data=f"send:{chat_id}")])
        
        teclado = InlineKeyboardMarkup(botones)
        msg_exito = "✨ ¡Minuta estructurada con éxito! ¿A qué grupo privado de alumnos deseas enviarla?"
        
        if is_query: await update_or_query.edit_message_text(msg_exito, reply_markup=teclado)
        else: await update_or_query.message.reply_text(msg_exito, reply_markup=teclado)
        
        if grupos_sin_permisos:
            id_prof = update_or_query.effective_user.id
            lista_formateada = ", ".join([f"<b>{g}</b>" for g in grupos_sin_permisos])
            await context.bot.send_message(
                chat_id=id_prof,
                text=f"🛑 <b>Notificación de Permisos Críticos:</b>\nEl Delegado Virtual <b>NO</b> posee accesos de administrador en: {lista_formateada}.\n\n⚠️ <i>Otorgue rango de Administrador en esas comunidades para permitir publicaciones.</i>",
                parse_mode="HTML"
            )


# --- MODIFICACIÓN 2: EXTRACTOR DE REGLAS DE CONVIVENCIA DESDE AUDIO ---
async def ejecutar_flujo_configurar_reglas(update, context, datos_audio_bytes, chat_id_grupo):
    prompt_reglas = (
        "Analiza el audio del profesor donde dicta las normas de convivencia del grupo académico.\n"
        "Estructura un reglamento interno universitario impecable, formal y numerado.\n"
        "Asegúrate de enfatizar la prohibición estricta de material obsceno, lenguaje ofensivo/grosero, spam, "
        "y aclara que el canal es de uso exclusivo para temas de la materia y comprobantes de pago.\n"
        "Devuelve exclusivamente el texto formateado en HTML limpio admisible por Telegram (<b>, <i>)."
    )
    try:
        respuesta = ai_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[types.Part.from_bytes(data=bytes(datos_audio_bytes), mime_type="audio/ogg"), prompt_reglas]
        )
        texto_reglas = sanitizar_html_telegram(respuesta.text)
        
        # Guardar en memoria local
        guardar_reglas(chat_id_grupo, texto_reglas)
        
        grupos = cargar_grupos()
        nombre_g = grupos.get(str(chat_id_grupo), {}).get("titulo", "Grupo Académico")
        
        # Publicar y fijar en el canal estudiantil de forma nativa
        mensaje_reglas = await context.bot.send_message(
            chat_id=int(chat_id_grupo),
            text=f"📜 <b>[REGLAMENTO INSTITUCIONAL - {nombre_g.upper()}]</b>\n\n{texto_reglas}\n\n"
                 f"⚠️ <i>Nota: Este grupo es auditado en tiempo real por el Delegado Virtual. "
                 f"Romper las reglas 3 veces en una ventana de una semana conlleva a suspensiones temporales de escritura (24h) o la expulsión definitiva del grupo.</i>",
            parse_mode="HTML"
        )
        
        await context.bot.pin_chat_message(chat_id=int(chat_id_grupo), message_id=mensaje_reglas.message_id)
        await update.message.reply_text(f"✅ ¡Reglamento procesado, publicado y fijado con éxito en el grupo <b>{nombre_g}</b>!", parse_mode="HTML")
    except Exception as e:
        logging.error(f"Error al configurar reglamento: {e}")
        await update.message.reply_text(f"❌ Error crítico al procesar el reglamento de convivencia con Gemini: {e}")


# =========================================================================
# DETECTOR DE CAPTURES POR VISIÓN INTELIGENTE (ANTIFRAUDE - OPCIÓN 1)
# =========================================================================
async def supervisar_grupo_fotos_pagos(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat = update.effective_chat
    user = update.effective_user
    
    evento = obtener_evento_activo(chat.id)
    if not evento or evento.get("estado") == "cerrado":
        return 

    fecha_limite = datetime.strptime(evento["fecha_limite"], "%Y-%m-%d %H:%M:%S")
    if datetime.now() > fecha_limite:
        cerrar_evento_en_json(chat.id)
        await update.message.reply_text("🚨 <b>Periodo cerrado:</b> El tiempo límite fijado por el docente para esta evaluación concluyó. Pago no procesado.")
        return

    foto_file = await update.message.photo[-1].get_file()
    bytes_foto = await foto_file.download_as_bytearray()
    
    imagen_part = types.Part.from_bytes(
        data=bytes(bytes_foto),
        mime_type="image/jpeg"
    )
    
    prompt = (
        "Analiza detalladamente este capture de pantalla de transferencia bancaria.\n"
        "Debes extraer obligatoriamente dos datos:\n"
        "1. El nombre y apellido de la persona que realiza o es dueña del pago.\n"
        "2. El número de referencia de la operación bancaria.\n"
        "Devuelve un formato JSON estructurado estricto con las siguientes llaves exactas:\n"
        '{"nombre": "string o null", "referencia": "string o null", "datos_completos": true/false}'
    )
    
    try:
        response = ai_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[imagen_part, prompt],
            config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        
        resultado = json.loads(response.text)
        
        if not resultado.get("datos_completos") or not resultado.get("nombre") or not resultado.get("referencia"):
            await update.message.reply_text(
                f"⚠️ El estudiante {user.mention_html()} ha enviado un comprobante, pero no se visualiza legiblemente su **Nombre/Apellido** o el **Número de Referencia**.\nPor favor vuelve a subir el capture completo.",
                parse_mode="HTML"
            )
            return

        nombre_detectado = resultado["nombre"].strip()
        ref_detectada = str(resultado["referencia"]).strip()
        evento_nombre = evento["evento_activo"]
        
        resultado_registro = registrar_pago_valido(chat.id, nombre_detectado, ref_detectada, evento_nombre)
        
        if resultado_registro == "duplicado_referencia":
            await update.message.reply_text(
                f"❌ El pago reportado por {user.mention_html()} ha sido **RECHAZADO**.\nEl número de referencia <code>{ref_detectada}</code> ya fue validado previamente.",
                parse_mode="HTML"
            )
        elif resultado_registro == "duplicado_estudiante":
            await update.message.reply_text(
                f"⚠️ Atención {user.mention_html()}, tu nombre ya figura en la lista de concilación aprobada para la presente recaudación.",
                parse_mode="HTML"
            )
        elif resultado_registro == "exito":
            pagos_totales = cargar_pagos().get(str(chat.id), {}).get(evento_nombre, [])
            contador_actual = len(pagos_totales)
            total_estudiantes = evento["estudiantes_objetivo"]
            
            await update.message.reply_text(
                f"✅ <b>Comprobante de Pago Validado</b>\n"
                f"• <b>Estudiante Extraído:</b> {nombre_detectado}\n"
                f"• <b>N° Referencia:</b> <code>{ref_detectada}</code>\n"
                f"• <b>Evaluación:</b> {evento_nombre}\n"
                f"• <b>Control de Aula:</b> {contador_actual}/{total_estudiantes} reportados con éxito.",
                parse_mode="HTML"
            )
            
            if contador_actual >= total_estudiantes:
                cerrar_evento_en_json(chat.id)
                await update.message.reply_text(
                    f"🎉 ¡Excelente! El 100% de la comunidad estudiantil de esta sección ({contador_actual}/{total_estudiantes}) ha completado su abono para <b>{evento_nombre}</b>. El flujo de registro queda oficialmente cerrado."
                )
                
                grupos = cargar_grupos()
                id_profesor = grupos.get(str(chat.id), {}).get("id_profesor")
                if id_profesor:
                    await context.bot.send_message(
                        chat_id=id_profesor,
                        text=f"🚀 <b>¡Meta de Recaudación Cumplida!</b> El 100% de los estudiantes de la sección <b>{grupos[str(chat.id)]['titulo']}</b> han completado sus aportes para <b>{evento_nombre}</b>. El registro se ha cerrado de forma automática.",
                        parse_mode="HTML"
                    )
                    
    except Exception as e:
        logging.error(f"Fallo en Visión Artificial: {e}")


# Tarea en Segundo Plano ejecutada al cumplirse el plazo límite
async def ejecutar_cierre_automatico(context: ContextTypes.DEFAULT_TYPE):
    job = context.job
    chat_id = job.data["chat_id"]
    evento_nombre = job.data["evento_nombre"]
    id_profesor = job.data["id_profesor"]
    
    evento = obtener_evento_activo(chat_id)
    if evento and evento["estado"] == "abierto" and evento["evento_activo"] == evento_nombre:
        cerrar_evento_en_json(chat_id)
        
        await context.bot.send_message(
            chat_id=chat_id,
            text=f"🚨 <b>Periodo de Recaudación Finalizado</b> 🚨\n\nLa fecha y hora límite para reportar los pagos de <b>{evento_nombre}</b> ha concluido.\nA partir de este momento, el Delegado Virtual ya no validará más comprobantes bancarios automáticamente.",
            parse_mode="HTML"
        )
        
        await context.bot.send_message(
            chat_id=id_profesor,
            text=f"📋 <b>Notificación de Cierre Temporal:</b> El plazo límite para recaudar dinero destinado a <b>{evento_nombre}</b> ha expirado. Las admisiones están cerradas.",
            parse_mode="HTML"
        )


# =========================================================================
# MODIFICACIÓN 2: MOTOR DISCIPLINARIO DE AUDITORÍA DE CONDUCTA EN VIVO
# =========================================================================
async def auditar_comportamiento_estudiante(update: Update, context: ContextTypes.DEFAULT_TYPE):
    chat = update.effective_chat
    user = update.effective_user
    
    if not chat or chat.type in ["private"] or not update.message or user.is_bot:
        return
        
    reglas_sistema = cargar_reglas()
    if str(chat.id) not in reglas_sistema:
        return  # No hay reglamento configurado para esta sección

    # Los profesores y admins del canal están exentos de penalizaciones
    try:
        miembro = await chat.get_member(user.id)
        if miembro.status in [ChatMember.ADMINISTRATOR, ChatMember.OWNER]:
            return
    except Exception:
        pass

    # Evitar interferencias con la Modificación 3 (Menciones directas para consultas)
    bot_username = context.bot.username
    if bot_username and f"@{bot_username}" in (update.message.text or ""):
        return

    contenido_evaluar = update.message.text or update.message.caption or ""
    if not contenido_evaluar:
        return

    prompt_auditor = (
        f"Actúas como un prefecto de disciplina universitario. Evalúa si el siguiente mensaje de un alumno "
        f"rompe las reglas del grupo académico.\n\n"
        f"REGLAMENTO DEL GRUPO:\n{reglas_sistema[str(chat.id)]}\n\n"
        f"MENSAJE DEL ALUMNO:\n'{contenido_evaluar}'\n\n"
        f"Devuelve estrictamente un JSON con este formato exacto:\n"
        f'{{"rompe_regla": true/false, "motivo": "Explicación breve de la falta en una línea"}}'
    )
    
    try:
        response = ai_client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[prompt_auditor],
            config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        veredicto = json.loads(response.text)
        
        if not veredicto.get("rompe_regla"):
            return 
            
        motivo_falta = veredicto.get("motivo", "Conducta inapropiada / Mensaje ajeno a la cátedra")
        ahora = datetime.now()
        str_chat = str(chat.id)
        str_user = str(user.id)
        
        comp_data = cargar_comportamiento()
        
        if str_chat not in comp_data: comp_data[str_chat] = {}
        if str_user not in comp_data[str_chat]:
            comp_data[str_chat][str_user] = {"infracciones": [], "nivel_sancion": 0}
            
        comp_data[str_chat][str_user]["infracciones"].append(ahora.strftime("%Y-%m-%d %H:%M:%S"))
        
        # Filtro de Ventana Móvil de 7 días (Strikes semanales)
        infracciones_validas = []
        for f_str in comp_data[str_chat][str_user]["infracciones"]:
            f_dt = datetime.strptime(f_str, "%Y-%m-%d %H:%M:%S")
            if (ahora - f_dt).days < 7:
                infracciones_validas.append(f_str)
                
        comp_data[str_chat][str_user]["infracciones"] = infracciones_validas
        strikes_semanales = len(infracciones_validas)
        
        if strikes_semanales < 3:
            # Advertencia e informe de saldo de strikes
            await update.message.reply_text(
                f"⚠️ <b>Advertencia de Convivencia</b>\n"
                f"Estudiante: {user.mention_html()}\n"
                f"Motivo: {html.escape(motivo_falta)}.\n"
                f"⚠️ Tienes <b>{strikes_semanales}/3</b> llamadas de atención esta semana. Al llegar a 3 serás suspendido temporalmente.",
                parse_mode="HTML"
            )
            guardar_comportamiento(comp_data)
        else:
            # Se alcanzaron 3 strikes en la semana -> Escalar Sanción y resetear buffer semanal
            comp_data[str_chat][str_user]["nivel_sancion"] += 1
            nivel = comp_data[str_chat][str_user]["nivel_sancion"]
            comp_data[str_chat][str_user]["infracciones"] = []
            guardar_comportamiento(comp_data)
            
            if nivel in [1, 2]:
                # MUTEO TEMPORAL POR 24 HORAS
                hora_desmuteo = ahora + timedelta(days=1)
                await context.bot.restrict_chat_member(
                    chat_id=chat.id,
                    user_id=user.id,
                    permissions=ChatPermissions(can_send_messages=False),
                    until_date=hora_desmuteo
                )
                await context.bot.send_message(
                    chat_id=chat.id,
                    text=f"🛑 <b>Estudiante Suspendido (Sanción Nivel {nivel})</b>\n\n"
                         f"El alumno {user.mention_html()} ha acumulado 3 infracciones reglamentarias en la misma semana.\n"
                         f"• <b>Acción:</b> Suspensión de permisos de escritura en el grupo (Mute).\n"
                         f"• <b>Duración:</b> 24 Horas.\n"
                         f"• <b>Razón de la última falta:</b> {html.escape(motivo_falta)}.",
                    parse_mode="HTML"
                )
            else:
                # TERCER CONTINGENCIA: BANEO Y EXPULSIÓN DEFINITIVA
                await context.bot.ban_chat_member(chat_id=chat.id, user_id=user.id)
                await context.bot.send_message(
                    chat_id=chat.id,
                    text=f"🚨 <b>Expulsión por Reincidencia Crítica</b> 🚨\n\n"
                         f"El estudiante {user.mention_html()} ha violado el reglamento de la sección por tercera vez consecutiva tras cumplir sus suspensiones temporales.\n"
                         f"• <b>Acción:</b> Expulsado del grupo de Telegram de forma permanente.",
                    parse_mode="HTML"
                )
    except Exception as e:
        logging.error(f"Error en el motor de convivencia activa: {e}")


# =========================================================================
# MODIFICACIÓN 3: INBOX DE PREGUNTAS AL DOCENTE CON CUPO DIARIO (MAX 10)
# =========================================================================
async def gestionar_consulta_estudiante(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    if not message or not message.text:
        return

    bot_username = context.bot.username
    tag_bot = f"@{bot_username}"

    if tag_bot not in message.text:
        return  # Si no invocan al bot, ignorar para que pase a los siguientes handlers

    chat_grupo = update.effective_chat
    user_estudiante = update.effective_user
    texto_original = message.text
    ahora_fecha = datetime.now().strftime("%Y-%m-%d")
    
    grupos = cargar_grupos()
    str_chat = str(chat_grupo.id)
    
    # Validar que el grupo esté indexado para ubicar al profesor dueño del aula
    if str_chat not in grupos:
        return

    id_profesor = grupos[str_chat].get("id_profesor")
    if not id_profesor:
        return

    consultas_data = cargar_consultas()
    if str_chat not in consultas_data:
        consultas_data[str_chat] = {"fecha": ahora_fecha, "contador": 0}
    
    # Reinicio del cupo al cambiar de día en el servidor
    if consultas_data[str_chat]["fecha"] != ahora_fecha:
        consultas_data[str_chat]["fecha"] = ahora_fecha
        consultas_data[str_chat]["contador"] = 0

    contador_actual = consultas_data[str_chat]["contador"]
    MAX_PREGUNTAS = 10

    if contador_actual >= MAX_PREGUNTAS:
        await message.reply_text(
            f"❌ <b>Cupo de Consultas Agotado</b>\n\n"
            f"Lo siento, esta sección ya ha alcanzado el límite de <b>{MAX_PREGUNTAS} preguntas diarias</b> permitidas para el profesor.\n"
            f"⏳ El cupo se restablecerá mañana de forma automática para garantizar la optimización de tiempos del docente.",
            parse_mode="HTML"
        )
        return

    # Limpiar etiqueta y aislar la duda académica
    pregunta_limpia = texto_original.replace(tag_bot, "").strip()
    if not pregunta_limpia:
        await message.reply_text("⚠️ ¡Hola! Para mandarle una duda al profesor, escribe tu pregunta detallada después de etiquetarme.")
        return

    # Consumir saldo de consultas
    consultas_data[str_chat]["contador"] += 1
    guardar_consultas(consultas_data)
    
    preguntas_restantes = MAX_PREGUNTAS - consultas_data[str_chat]["contador"]
    nombre_grupo = grupos[str_chat].get("titulo", "Grupo Académico")
    nombre_alumno = user_estudiante.full_name

    # Enviar al Profesor en Privado aplicando filtros sanitizadores HTML
    texto_para_profesor = (
        f"📥 <b>[NUEVA CONSULTA ESTUDIANTIL]</b>\n\n"
        f"• <b>Grupo:</b> {html.escape(nombre_grupo)}\n"
        f"• <b>Estudiante:</b> {html.escape(nombre_alumno)}\n"
        f"• <b>Pregunta:</b> <i>\"{html.escape(pregunta_limpia)}\"</i>\n\n"
        f"⚠️ <i>Nota: Puedes coordinar la respuesta directamente ingresando a la sección o contactando al alumno en privado.</i>"
    )

    try:
        await context.bot.send_message(
            chat_id=int(id_profesor),
            text=texto_para_profesor,
            parse_mode="HTML"
        )
        
        # Responder con éxito informando el cupo restante de la sección
        await message.reply_text(
            f"📥 ✨ <b>¡Consulta transmitida con éxito!</b>\n"
            f"He enviado tu duda de forma directa al privado del profesor.\n\n"
            f"📊 <b>Estado del cupo diario de la sección:</b>\n"
            f"• Solicitudes realizadas hoy: {consultas_data[str_chat]['contador']}/{MAX_PREGUNTAS}\n"
            f"• Consultas disponibles restantes: <b>{preguntas_restantes}</b>",
            parse_mode="HTML"
        )
    except Exception as e:
        logging.error(f"Error al enviar consulta al profesor: {e}")
        await message.reply_text(
            "❌ <b>Error de Enrutamiento</b>\n"
            "No pude contactar al profesor en su buzón privado. Verifique que tenga el bot iniciado.",
            parse_mode="HTML"
        )


# =========================================================================
# ENRUTADOR CENTRAL DE INTERACCIONES INTERACTIVAS (CALLBACK QUERIES)
# =========================================================================
async def procesar_clicks_botones(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    data = query.data
    id_usuario = update.effective_user.id

    if data == "auth_si":
        context.user_data['esperando_clave'] = True
        await query.edit_message_text("Por favor, ingresa la contraseña de acceso:")
        return
    elif data == "auth_no":
        await query.edit_message_text("Acceso denegado, no eres un profesor.")
        return

    elif data.startswith("g_aceptar:"):
        chat_id_grupo = data.split(":")[1]
        titulo_grupo = context.bot_data.get(f"titulo_pendiente_{chat_id_grupo}", "Grupo Confirmado")
        guardar_grupo(chat_id_grupo, titulo_grupo, id_usuario)
        await query.edit_message_text(f"✅ El grupo <b>{titulo_grupo}</b> ha sido aceptado y enlazado exclusivamente a su panel de control.", parse_mode="HTML")
        return
    elif data.startswith("g_rechazar:"):
        chat_id_grupo = data.split(":")[1]
        titulo_grupo = context.bot_data.get(f"titulo_pendiente_{chat_id_grupo}", "Grupo Descartado")
        eliminar_grupo(chat_id_grupo)
        try:
            await query.edit_message_text(f"🚪 Solicitud rechazada. Saliendo del entorno estudiantil de <b>{titulo_grupo}</b>...", parse_mode="HTML")
            await context.bot.leave_chat(chat_id=chat_id_grupo)
        except Exception:
            pass
        return

    elif data.startswith("recaudar_g:"):
        chat_id_grupo = data.split(":")[1]
        context.user_data['esperando_recaudacion_datos'] = chat_id_grupo
        grupos = cargar_grupos()
        nombre_g = grupos.get(str(chat_id_grupo), {}).get("titulo", "Grupo")
        await query.edit_message_text(
            f"💰 Configurando recaudación para: <b>{nombre_g}</b>.\n\n"
            f"Por favor, envíame un mensaje de texto indicando el nombre del evento y la fecha/hora límite exacta separados obligatoriamente por una barra vertical:\n"
            f"<code>Nombre del Evento | AAAA-MM-DD HH:MM</code>\n\n"
            f"Ejemplo real:\n<code>Examen Unidad 1 | 2026-06-30 16:00</code>",
            parse_mode="HTML"
        )
        return

    # --- MODIFICACIÓN 2: SELECCIÓN INTERACTIVA PARA CONFIGURAR REGLAS ---
    elif data.startswith("config_reglas_g:"):
        chat_id_grupo = data.split(":")[1]
        context.user_data['esperando_audio_reglas'] = chat_id_grupo
        grupos = cargar_grupos()
        nombre_g = grupos.get(str(chat_id_grupo), {}).get("titulo", "Grupo")
        await query.edit_message_text(
            f"📜 <b>Configurando Reglamento para: {nombre_g}</b>\n\n"
            f"Por favor, envíame a continuación la <b>nota de voz</b> dictando las normas de convivencia del aula.\n\n"
            f"El Delegado Virtual procesará el audio con Inteligencia Artificial, estructurará el marco legal académico, lo publicará y lo fijará en el chat automáticamente.",
            parse_mode="HTML"
        )
        return

    elif data.startswith("env_mat:"):
        chat_id_destino = data.split(":")[1]
        material = context.user_data.get('material_pendiente')
        grupos = cargar_grupos()
        nombre_grupo = grupos.get(str(chat_id_destino), {}).get("titulo", "el grupo")

        if not material:
            await query.edit_message_text("⚠️ El material expiró en memoria o ya fue enviado. Reenvía el archivo original.")
            return

        try:
            if material["tipo"] == "documento":
                await context.bot.send_document(
                    chat_id=int(chat_id_destino),
                    document=material["file_id"],
                    caption=f"📚 <b>NUEVO MATERIAL DE APOYO</b>\n\nCompartido por el Profesor.\n<i>{material['caption']}</i>",
                    parse_mode="HTML"
                )
            elif material["tipo"] == "enlace":
                await context.bot.send_message(
                    chat_id=int(chat_id_destino),
                    text=f"🔗 <b>RECURSO O BIBLIOGRAFÍA COMPARTIDA</b>\n\nEl profesor ha compartido el siguiente enlace de interés:\n\n{material['contenido']}",
                    parse_mode="HTML"
                )
            await query.edit_message_text(f"✅ ¡Éxito! El material fue enviado a la sección de alumnos: <b>{nombre_grupo}</b>", parse_mode="HTML")
            context.user_data['material_pendiente'] = None
        except Exception as e:
            await query.edit_message_text(f"❌ Error crítico de despacho: {e}")
        return

    elif data == "cancelar_envio_mat":
        context.user_data['material_pendiente'] = None
        await query.edit_message_text("❌ Envío de material académico cancelado.")
        return

    elif data.startswith("audio_opt:"):
        accion = data.split(":")[1]
        audio_bytes_list = context.user_data.get('audio_bytes_pendientes')
        
        if not audio_bytes_list:
            await query.edit_message_text("❌ El archivo de audio caducó en la memoria volátil. Envíalo de nuevo.")
            return
            
        datos_audio_bytes = bytes(audio_bytes_list)
        
        if accion == "cancelar":
            context.user_data['audio_bytes_pendientes'] = None
            await query.edit_message_text("❌ Procesamiento de clase magistral cancelado.")
            return
             
        await query.edit_message_text("🤖 Conectando y analizando extensamente el flujo de audio con Gemini... por favor espera.")
        
        audio_part = types.Part.from_bytes(data=datos_audio_bytes, mime_type="audio/ogg")
        
        if accion == "mensaje":
            grupos_totales = cargar_grupos()
            grupos_activos = {cid: info["titulo"] for cid, info in grupos_totales.items() if info.get("id_profesor") == id_usuario}
            await ejecutar_flujo_audio_corto(query, context, datos_audio_bytes, grupos_activos)
            
        elif accion == "informe":
            prompt = (
                "Actúa como un redactor académico universitario de alto nivel. Desarrolla un informe o trabajo de estudio extenso, "
                "exhaustivo y sumamente completo basado enteramente en la información técnica descrita en este audio.\n"
                "Organiza la información separando obligatoriamente las secciones principales usando '# ' para el Título Principal y '## ' para los Subtítulos.\n"
                "El cuerpo restante debe ser redactado con prosa fluida, formal y en párrafos íntegros sin viñetas informales."
            )
            try:
                respuesta_ia = ai_client.models.generate_content(model='gemini-2.5-flash', contents=[audio_part, prompt])
                
                if respuesta_ia and respuesta_ia.text:
                    ruta_word = crear_documento_academico(respuesta_ia.text, "Informe_De_Estudio_Academico.docx")
                    
                    with open(ruta_word, "rb") as f:
                        await context.bot.send_document(
                            chat_id=id_usuario,
                            document=f,
                            filename="Informe_De_Estudio_Academico.docx",
                            caption="📄 <b>¡Informe Académico Generado!</b>\n\nAquí tienes el documento structured milimétricamente en fuente <i>Times New Roman</i>, interlineado 1.5, texto justificado y sangrías de 1.5cm al inicio de cada párrafo listo para su distribución.",
                            parse_mode="HTML"
                        )
                    if os.path.exists(ruta_word):
                        os.remove(ruta_word)
                else:
                    await context.bot.send_message(chat_id=id_usuario, text="⚠️ Imposible extraer contenido didáctico de la nota de voz.")
            except Exception as e:
                await context.bot.send_message(chat_id=id_usuario, text=f"❌ Fallo crítico en el motor documentador: {e}")
                
        context.user_data['audio_bytes_pendientes'] = None
        return

    elif data.startswith("send:"):
        chat_id_destino = data.split(":")[1]
        minuta = context.user_data.get('minuta_pendiente')
        
        if not minuta:
            await query.edit_message_text("⚠️ Error: La minuta caducó en memoria. Envía el audio nuevamente.")
            return
        
        grupos = cargar_grupos()
        info_grupo = grupos.get(chat_id_destino, {})
        nombre_grupo = info_grupo.get("titulo", "el grupo seleccionado")
        
        if info_grupo.get("id_profesor") != id_usuario:
            await query.edit_message_text("❌ Violación de permisos: No posees derechos administrativos sobre este grupo.")
            return
            
        try:
            miembro_bot = await context.bot.get_chat_member(chat_id=chat_id_destino, user_id=context.bot.id)
            if miembro_bot.status not in [ChatMember.ADMINISTRATOR, ChatMember.OWNER]:
                await query.edit_message_text(
                    f"❌ <b>Error de Publicación:</b> El bot carece de permisos de administrador en <b>{nombre_grupo}</b>.",
                    parse_mode="HTML"
                )
                return
        except Exception:
            await query.edit_message_text(f"❌ <b>Error de canal:</b> No se pudo establecer conexión con <b>{nombre_grupo}</b>.", parse_mode="HTML")
            return
            
        await query.edit_message_text("🚀 Transmitiendo información al canal estudiantil...")
        
        try:
            await context.bot.send_message(chat_id=chat_id_destino, text=minuta, parse_mode="HTML")
            await context.bot.send_message(
                chat_id=id_usuario, 
                text=f"✨ Minuta enviada con éxito al grupo <b>{nombre_grupo}</b>",
                parse_mode="HTML"
            )
            context.user_data['minuta_pendiente'] = None
        except Exception as e:
            await context.bot.send_message(chat_id=id_usuario, text=f"❌ Error al transmitir los datos al grupo: {e}")
        return
# =========================================================================
# CONTROLADOR Y MOTOR DE ARRANQUE CENTRAL
# =========================================================================
